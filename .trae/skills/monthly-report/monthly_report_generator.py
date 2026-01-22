#!/usr/bin/env python3
"""
GitLab月报生成工具
自动解析GitLab导出数据，生成结构化的月度工作报告
"""

import csv
import json
import re
from datetime import datetime, timedelta
from collections import defaultdict, Counter
from pathlib import Path
from typing import Dict, List, Any, Optional
import argparse
import sys

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class CommitData:
    """Git提交数据类"""
    def __init__(self, date: str, title: str, description: str, author: str, email: str):
        self.date = date
        self.title = title
        self.description = description
        self.author = author
        self.email = email
        self.category = self._categorize_commit()
        self.features = self._extract_features()
    
    def _categorize_commit(self) -> str:
        """根据标题和描述分类提交"""
        title_lower = self.title.lower()
        desc_lower = self.description.lower()
        combined = f"{title_lower} {desc_lower}"
        
        categories = {
            '功能开发': ['feat', '功能', '添加', '新增', '实现', '开发'],
            '问题修复': ['fix', '修复', 'bug', '错误', '修复'],
            '样式优化': ['style', '样式', 'ui', '布局', '界面', '设计'],
            '代码重构': ['refactor', '重构', '重构', '优化'],
            '文档更新': ['docs', '文档', '注释', '说明'],
            '测试相关': ['test', '测试', '测试用例'],
            '配置变更': ['config', '配置', '设置', '环境'],
            '资源管理': ['assets', '资源', '图片', '文件']
        }
        
        for category, keywords in categories.items():
            if any(keyword in combined for keyword in keywords):
                return category
        
        return '其他'
    
    def _extract_features(self) -> List[str]:
        """提取功能特征"""
        features = []
        
        # 提取括号内的功能模块
        bracket_pattern = r'\(([^)]+)\)'
        brackets = re.findall(bracket_pattern, self.title)
        features.extend(brackets)
        
        # 提取冒号后的主要功能
        colon_pattern = r'[:：]([^,，\n]+)'
        colons = re.findall(colon_pattern, self.title)
        for colon in colons:
            features.append(colon.strip())
        
        return features[:3]  # 最多保留3个特征


class MonthlyReportGenerator:
    """月度报告生成器"""
    
    def __init__(self, csv_file_path: str):
        self.csv_file_path = Path(csv_file_path)
        self.commits: List[CommitData] = []
        self.month_data = defaultdict(list)
        self.stats = {
            'total_commits': 0,
            'categories': Counter(),
            'authors': Counter(),
            'daily_activity': defaultdict(int),
            'features': Counter()
        }
    
    def parse_csv(self) -> bool:
        """解析CSV文件"""
        try:
            with open(self.csv_file_path, 'r', encoding='utf-8-sig') as file:
                reader = csv.DictReader(file)
                
                for row in reader:
                    if not row.get('时间') or row.get('时间') == '时间':
                        continue
                    
                    commit = CommitData(
                        date=row.get('时间', '').strip('"'),
                        title=row.get('标题', '').strip('"'),
                        description=row.get('描述', '').strip('"'),
                        author=row.get('作者', '').strip('"'),
                        email=row.get('邮箱', '').strip('"<>')
                    )
                    
                    self.commits.append(commit)
                    self._process_commit(commit)
            
            self.stats['total_commits'] = len(self.commits)
            return True
            
        except Exception as e:
            print(f"解析CSV文件时出错: {e}")
            return False
    
    def _process_commit(self, commit: CommitData):
        """处理提交数据"""
        # 按日期分组
        try:
            date_obj = datetime.strptime(commit.date, '%Y-%m-%d')
            date_key = date_obj.strftime('%Y-%m-%d')
            self.month_data[date_key].append(commit)
            self.stats['daily_activity'][date_key] += 1
        except ValueError:
            # 处理其他日期格式
            self.month_data[commit.date].append(commit)
        
        # 统计分类
        self.stats['categories'][commit.category] += 1
        
        # 统计作者
        self.stats['authors'][commit.author] += 1
        
        # 统计功能特征
        for feature in commit.features:
            if feature:
                self.stats['features'][feature] += 1
    
    def generate_markdown_report(self, output_path: str = None) -> str:
        """生成Markdown格式报告"""
        if not output_path:
            output_path = f"月度工作报告_{datetime.now().strftime('%Y%m')}.md"
        
        report_content = self._build_report_content()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_content)
        
        return output_path
    
    def generate_word_report(self, output_path: str = None) -> str:
        """生成Word格式报告"""
        if not DOCX_AVAILABLE:
            print("警告: python-docx未安装，无法生成Word文档")
            return None
        
        if not output_path:
            output_path = f"月度工作报告_{datetime.now().strftime('%Y%m')}.docx"
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(f'{datetime.now().strftime("%Y年%m月")} 开发工作报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本统计信息
        doc.add_heading('本月工作统计', level=1)
        stats_para = doc.add_paragraph()
        stats_para.add_run(f"• 总提交次数: {self.stats['total_commits']}\n")
        stats_para.add_run(f"• 活跃天数: {len(self.month_data)}\n")
        stats_para.add_run(f"• 主要分类: {', '.join([f'{k}({v})' for k, v in self.stats['categories'].most_common(3)])}\n")
        
        # 每日工作记录
        doc.add_heading('每日工作记录', level=1)
        
        # 获取唯一作者列表，如果只有一个作者则显示
        unique_authors = list(set(commit.author for commit in self.commits))
        if len(unique_authors) == 1:
            author_para = doc.add_paragraph()
            author_para.add_run(f"• 主要开发者: {unique_authors[0]}\n")
        
        for date, commits in sorted(self.month_data.items()):
            doc.add_heading(date, level=2)
            
            for commit in commits:
                # 简化的提交记录格式
                commit_para = doc.add_paragraph()
                commit_para.add_run(f"• {commit.title}").bold = True
                
                # 只在重要变更时显示描述
                if commit.description and any(keyword in commit.title.lower() for keyword in ['feat', '新增', '实现', '修复']):
                    desc_lines = commit.description.split('\n')[:2]  # 只显示前两行描述
                    for line in desc_lines:
                        clean_line = line.strip().replace('- ', '')  # 移除多余的 "- " 符号
                        if clean_line:
                            commit_para.add_run(f"\n  - {clean_line}")
        
        # 分类统计
        doc.add_heading('工作分类统计', level=1)
        for category, count in self.stats['categories'].most_common():
            doc.add_paragraph(f"• {category}: {count} 次")
        
        # 功能模块统计
        if self.stats['features']:
            doc.add_heading('主要功能模块', level=1)
            for feature, count in self.stats['features'].most_common(10):
                doc.add_paragraph(f"• {feature}: {count} 次")
        
        doc.save(output_path)
        return output_path
    
    def _build_report_content(self) -> str:
        """构建报告内容"""
        content = []
        
        # 标题
        content.append(f"# {datetime.now().strftime('%Y年%m月')} 开发工作报告\n")
        
        # 基本统计
        content.append("## 📊 本月工作统计\n")
        content.append(f"- **总提交次数**: {self.stats['total_commits']} 次")
        content.append(f"- **活跃天数**: {len(self.month_data)} 天")
        content.append(f"- **主要工作分类**: {', '.join([f'{k}({v})' for k, v in self.stats['categories'].most_common(3)])}\n")
        
        # 每日工作记录
        content.append("## 📅 每日工作记录\n")
        
        # 获取唯一作者列表，如果只有一个作者则显示
        unique_authors = list(set(commit.author for commit in self.commits))
        if len(unique_authors) == 1:
            content.append(f"**主要开发者**: {unique_authors[0]}\n")
        
        for date, commits in sorted(self.month_data.items()):
            content.append(f"### {date}\n")
            
            for commit in commits:
                # 简化的提交记录格式
                content.append(f"- **{commit.title}**")
                
                # 只在重要变更时显示描述
                if commit.description and any(keyword in commit.title.lower() for keyword in ['feat', '新增', '实现', '修复']):
                    desc_lines = commit.description.split('\n')[:2]  # 只显示前两行描述
                    for line in desc_lines:
                        clean_line = line.strip().replace('- ', '')  # 移除多余的 "- " 符号
                        if clean_line:
                            content.append(f"  - {clean_line}")
                
                content.append("")  # 空行分隔
        
        # 工作分类统计
        content.append("## 🏷️ 工作分类统计\n")
        for category, count in self.stats['categories'].most_common():
            content.append(f"- **{category}**: {count} 次")
        
        # 功能模块统计
        if self.stats['features']:
            content.append("\n## 🔧 主要功能模块\n")
            for feature, count in self.stats['features'].most_common(10):
                content.append(f"- **{feature}**: {count} 次")
        
        # 工作总结
        content.append("\n## 📈 工作总结\n")
        content.append(self._generate_summary())
        
        return '\n'.join(content)
    
    def _generate_summary(self) -> str:
        """生成工作总结"""
        summary = []
        
        # 统计最活跃的日期
        if self.stats['daily_activity']:
            most_active_day = max(self.stats['daily_activity'].items(), key=lambda x: x[1])
            summary.append(f"- **最活跃工作日**: {most_active_day[0]} ({most_active_day[1]} 次提交)")
        
        # 统计主要工作类型
        if self.stats['categories']:
            main_work_type = self.stats['categories'].most_common(1)[0]
            summary.append(f"- **主要工作类型**: {main_work_type[0]} ({main_work_type[1]} 次)")
        
        # 统计主要功能模块
        if self.stats['features']:
            top_features = self.stats['features'].most_common(3)
            summary.append(f"- **重点功能模块**: {', '.join([f'{k}({v})' for k, v in top_features])}")
        
        return '\n'.join(summary)
    
    def export_json(self, output_path: str = None) -> str:
        """导出JSON格式数据"""
        if not output_path:
            output_path = f"commit_data_{datetime.now().strftime('%Y%m')}.json"
        
        export_data = {
            'month': datetime.now().strftime('%Y-%m'),
            'total_commits': self.stats['total_commits'],
            'daily_data': {},
            'categories': dict(self.stats['categories']),
            'authors': dict(self.stats['authors']),
            'features': dict(self.stats['features'])
        }
        
        for date, commits in self.month_data.items():
            export_data['daily_data'][date] = [
                {
                    'title': commit.title,
                    'description': commit.description,
                    'category': commit.category,
                    'author': commit.author,
                    'features': commit.features
                }
                for commit in commits
            ]
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)
        
        return output_path


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='GitLab月报生成工具')
    parser.add_argument('csv_file', help='GitLab导出的CSV文件路径')
    parser.add_argument('--output', '-o', help='输出文件路径')
    parser.add_argument('--format', '-f', choices=['md', 'word', 'json', 'all'], 
                       default='md', help='输出格式 (默认: md)')
    parser.add_argument('--install-deps', action='store_true', 
                       help='安装必要的Python依赖包')
    
    args = parser.parse_args()
    
    # 安装依赖
    if args.install_deps:
        print("正在安装必要的依赖包...")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 
                             'python-docx', 'markdown'])
        print("依赖包安装完成！")
        return
    
    # 检查输入文件
    if not Path(args.csv_file).exists():
        print(f"错误: CSV文件 '{args.csv_file}' 不存在")
        return
    
    # 创建生成器
    generator = MonthlyReportGenerator(args.csv_file)
    
    # 解析数据
    print("正在解析CSV文件...")
    if not generator.parse_csv():
        print("解析CSV文件失败")
        return
    
    print(f"成功解析 {generator.stats['total_commits']} 条提交记录")
    
    # 生成报告
    output_files = []
    
    if args.format in ['md', 'all']:
        md_file = generator.generate_markdown_report(args.output)
        output_files.append(md_file)
        print(f"Markdown报告已生成: {md_file}")
    
    if args.format in ['word', 'all']:
        word_file = generator.generate_word_report(args.output.replace('.md', '.docx') if args.output else None)
        if word_file:
            output_files.append(word_file)
            print(f"Word报告已生成: {word_file}")
        else:
            print("Word报告生成失败 (需要安装python-docx)")
    
    if args.format == 'json':
        json_file = generator.export_json(args.output.replace('.md', '.json') if args.output else None)
        output_files.append(json_file)
        print(f"JSON数据已导出: {json_file}")
    
    print(f"\n报告生成完成！共生成 {len(output_files)} 个文件")


if __name__ == '__main__':
    main()